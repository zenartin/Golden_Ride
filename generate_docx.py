from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    return heading

doc = Document()

# Title
title = doc.add_heading('Golden Ride - Deployment and Pricing Guide', 0)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

doc.add_paragraph("This document outlines the deployment strategy for the Golden Ride platform and provides a detailed breakdown of the costs involved.")

# Section 1
add_heading(doc, '1. Developer Account Costs', 1)
doc.add_paragraph("To publish the User App and Driver App on the official stores, you must create developer accounts.")

table1 = doc.add_table(rows=1, cols=4)
table1.style = 'Table Grid'
hdr_cells = table1.rows[0].cells
hdr_cells[0].text = 'Platform'
hdr_cells[1].text = 'Cost'
hdr_cells[2].text = 'Billing Cycle'
hdr_cells[3].text = 'Requirements'

row_cells = table1.add_row().cells
row_cells[0].text = 'Google Play Store'
row_cells[1].text = '$25'
row_cells[2].text = 'One-time fee'
row_cells[3].text = 'A Google account and a valid credit card.'

row_cells = table1.add_row().cells
row_cells[0].text = 'Apple App Store'
row_cells[1].text = '$99'
row_cells[2].text = 'Yearly'
row_cells[3].text = 'An Apple ID. (Registered as an Individual to avoid the D-U-N-S Number requirement).'

p = doc.add_paragraph()
runner = p.add_run("IMPORTANT: ")
runner.bold = True
runner.font.color.rgb = RGBColor(255, 0, 0)
p.add_run("Registering the Apple Developer Account as an Individual is highly recommended as it skips the lengthy business verification process (D-U-N-S Number). Approval is generally much faster.")

# Section 2
add_heading(doc, '2. Infrastructure Hosting Strategy (AWS)', 1)
doc.add_paragraph("You asked if the Backend and the Admin Website can be hosted on a single server to save costs. Yes, absolutely!")
doc.add_paragraph("Both the Python (FastAPI) Backend and the Admin Dashboard (React/Vite) can be hosted on a single Amazon Web Services (AWS) EC2 instance.")

add_heading(doc, 'How it works:', 2)
doc.add_paragraph("1. The Backend runs as a service on the AWS EC2 server.")
doc.add_paragraph("2. The Admin Website is compiled into static HTML/JS files and served by a lightweight proxy like Nginx running on the exact same AWS server.")
doc.add_paragraph("3. The Database (PostgreSQL) can also run on this same server for the initial launch phase to save money.")

table2 = doc.add_table(rows=1, cols=4)
table2.style = 'Table Grid'
hdr2_cells = table2.rows[0].cells
hdr2_cells[0].text = 'Service'
hdr2_cells[1].text = 'Recommended Provider'
hdr2_cells[2].text = 'Cost Estimate'
hdr2_cells[3].text = 'Purpose'

row2_cells = table2.add_row().cells
row2_cells[0].text = 'Combined Server'
row2_cells[1].text = 'AWS (Amazon Web Services) - EC2 t3.micro'
row2_cells[2].text = '~$8 to $10 / month'
row2_cells[3].text = 'Hosts the Backend API, PostgreSQL Database, and the Admin Dashboard.'

row2_cells = table2.add_row().cells
row2_cells[0].text = 'Domain Name'
row2_cells[1].text = 'AWS Route 53 / GoDaddy'
row2_cells[2].text = '~$12 to $15 / year'
row2_cells[3].text = 'A custom domain (e.g., goldenrideapp.com).'

# Section 3
add_heading(doc, '3. Third-Party API Costs', 1)
doc.add_paragraph("The platform relies on a few external services. Both offer generous free tiers that usually cover all initial launch usage.")
doc.add_paragraph("- Google Maps API: Gives you a $200 free credit every month. Unless you are doing thousands of rides per month, your map routing and location services will be virtually $0.")
doc.add_paragraph("- Razorpay: No monthly fees. They charge a standard percentage fee (usually ~2%) only when a successful transaction occurs.")

# Section 4
add_heading(doc, '4. Total Cost Summary to Present to Client', 1)

add_heading(doc, 'Upfront & Yearly Costs:', 2)
doc.add_paragraph("- Google Play Console License: $25 (One-time)")
doc.add_paragraph("- Apple Developer Program (Individual): $99 (Per Year)")
doc.add_paragraph("- Domain Name: ~$15 (Per Year)")
p2 = doc.add_paragraph()
p2.add_run("Total Initial Setup Cost: ~$139").bold = True

add_heading(doc, 'Ongoing Monthly Costs:', 2)
doc.add_paragraph("- Cloud Hosting (AWS EC2 - Backend, Database, Admin): ~$10.00 / month")
doc.add_paragraph("- Maps API: $0 (Covered by $200/mo free tier)")
doc.add_paragraph("- Payment Gateway: $0 monthly (Transaction fees apply per ride)")
p3 = doc.add_paragraph()
p3.add_run("Total Fixed Monthly Cost: ~$10 / month").bold = True

doc.add_paragraph("\nNote: This single-server strategy on AWS keeps the launch costs incredibly lean. Once the platform grows, AWS makes it very easy to scale up and separate the Database and Backend into their own larger services.")

doc.save(r'C:\Users\nanda\.gemini\antigravity\brain\e0b2eccd-98d7-4ae1-aa59-cfdeeee3d8c3\Deployment_and_Pricing.docx')
print("DOCX created successfully.")
